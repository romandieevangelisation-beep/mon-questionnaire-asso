import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import json
import re
from supabase import create_client
from cryptography.fernet import Fernet

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="Espace Clinique - YSQ-L3 Intégral", layout="wide")

@st.cache_resource
def init_connection():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return create_client(url, key)
    except: return None
supabase = init_connection()

# --- 2. SÉCURITÉ & DONNÉES ---
def get_cipher(): return Fernet(st.secrets["ENCRYPTION_KEY"].encode())

def decrypt_reponses(encrypted_json):
    """Lecture hybride : tente de déchiffrer, sinon lit en clair"""
    if not encrypted_json: return {}
    try:
        return json.loads(get_cipher().decrypt(encrypted_json.encode()).decode())
    except:
        try: return json.loads(encrypted_json)
        except: return {}

def save_patient_data(nom, email, reponses_dict):
    if not supabase: return False
    data = {"nom": nom, "email": email, "reponses_json": json.dumps(reponses_dict), "created_at": datetime.now().isoformat()}
    try:
        try: data["reponses_json"] = get_cipher().encrypt(json.dumps(reponses_dict).encode()).decode()
        except: pass
        supabase.table("patients_ysq").insert(data).execute()
        return True
    except: return False

def load_all_patients():
    if not supabase: return pd.DataFrame()
    try:
        response = supabase.table("patients_ysq").select("*").order("created_at", desc=True).execute()
        return pd.DataFrame(response.data)
    except: return pd.DataFrame()

def delete_patient(patient_id):
    if not supabase: return False
    try:
        supabase.table("patients_ysq").delete().eq("id", patient_id).execute()
        return True
    except: return False

# --- 3. BASE DE CONNAISSANCES COMPLETE (AVEC VERSETS ENTIERS) ---
DATA_SCHEMAS = {
    "ED": {
        "titre": "Carence Affective",
        "slogan": "« Personne ne me considère ou ne m'aime vraiment »",
        "clinique_expert": "Ce schéma signale un vide émotionnel précoce. Le patient a intégré la croyance que ses besoins de chaleur, d'empathie et de protection ne seront jamais validés par autrui. Il y a souvent une 'alexithymie' (difficulté à nommer ses besoins) par résignation acquise.",
        "theologie_expert": "Le mensonge racine est l'orphelinat spirituel. La guérison passe par la doctrine de l'Adoption (Romains 8:15). Dieu n'est pas un observateur distant mais un Père qui s'incline pour nourrir (Osée 11:4). Le travail consiste à réapprendre à 'recevoir' sans dissociation.",
        "origines": ["Le soignant ne répondait pas aux besoins d'affection ou de protection.", "Parents froids ou absents émotionnellement.", "L'enfant n'a pas été 'vu' ou compris dans sa singularité."],
        "symptomes": ["Sentiment chronique de vide ou de solitude.", "Considérer ses propres besoins comme sans importance.", "Devenir dépendant ou au contraire contre-dépendant (froid).", "Ressentiment envers les autres qui 'ne donnent pas assez'."],
        "mecanisme_titre": "Les 3 Formes de Privation",
        "mecanisme_texte": "1. Privation d'Attention (manque de chaleur).\n2. Privation d'Empathie (manque d'écoute).\n3. Privation de Protection (manque de guidance).",
        "actions_therapeute": ["Soyez à l'écoute de vos besoins d'attention et de protection.", "Exprimez vos besoins de manière assertive ('J'ai besoin de...').", "Évitez les partenaires froids qui recréent la privation familière."],
        "action_pastorale": "Tenez un journal de vos besoins devant Dieu. Osez demander une petite chose simple à un proche sans vous excuser, comme un acte de foi que Dieu vous aime et que vous n'avez pas besoin de le mériter.",
        "verset": "Psaume 27:10 - « Car mon père et ma mère m'abandonnent, mais l'Éternel me recueillera. »"
    },
    "AB": {
        "titre": "Abandon / Instabilité",
        "slogan": "« Ne me quitte pas »",
        "clinique_expert": "Perception de l'instabilité fondamentale des liens. Le patient vit dans l'hypervigilance de la perte, alternant entre agrippement anxieux et évitement préventif. La permanence de l'objet (l'autre) n'est pas acquise émotionnellement.",
        "theologie_expert": "L'antidote est la théologie de l'Alliance (Berit). Contrairement aux alliances humaines brisées, l'Alliance divine est unilatérale et irrévocable, fondée sur la fidélité de Dieu et non la performance humaine. Dieu est le Rocher (stabilité ontologique).",
        "origines": ["Décès d'un parent ou départ du foyer dans l'enfance.", "Soignant instable (dépression, alcool) ou imprévisible.", "Surprotection familiale rendant la séparation angoissante."],
        "symptomes": ["S'accrocher aux gens par peur (agrippement).", "Jalousie excessive et possessivité.", "Accusations injustifiées d'infidélité ou d'abandon."],
        "mecanisme_titre": "Le Cycle de l'Abandon",
        "mecanisme_texte": "1. Anxiété (recherche éperdue). 2. Colère/Désespoir (protestation). 3. Détachement (repli sur soi). Ce cycle de l'enfance se répète dans les relations adultes.",
        "actions_therapeute": ["Repérez votre tendance à dramatiser les séparations.", "Apprenez à vous apaiser seul(e) quand l'autre est absent.", "Évitez les partenaires instables ou ambivalents."],
        "action_pastorale": "Pratiquez la 'Solitude Habitée'. Passez 15 min seul(e) en visualisant la présence de Dieu. Rappelez-vous : 'Je ressens de la peur, mais je ne suis pas en danger réel'.",
        "verset": "Hébreux 13:5 - « Je ne te délaisserai point, et je ne t'abandonnerai point. »"
    },
    "MA": {
        "titre": "Méfiance / Abus",
        "slogan": "« Le monde est dangereux »",
        "clinique_expert": "Attente que l'autre va nuire, manipuler ou trahir. Le patient projette une intentionnalité malveillante sur autrui. C'est un schéma de survie traumatique où la confiance est vécue comme une mise en danger.",
        "theologie_expert": "Le monde est déchu, mais Dieu est le Refuge (Mahseh). La guérison demande de renoncer à l'auto-protection cynique pour accepter la protection de Dieu. C'est le passage de la suspicion (peur) au discernement (sagesse) sous le regard de Dieu.",
        "origines": ["Abus physique, sexuel ou verbal dans l'enfance.", "Famille humiliante, sadique ou punitive.", "Trahi ou manipulé par une figure de confiance."],
        "symptomes": ["Hypervigilance ('scanner' les menaces).", "Tests de loyauté envers les autres.", "Attaquer avant d'être attaqué."],
        "mecanisme_titre": "Types d'Abus & Méfiance",
        "mecanisme_texte": "Le schéma naît souvent d'abus physiques, sexuels ou verbaux. La personne reste en mode 'survie', s'attendant à ce que toute gentillesse cache un piège.",
        "actions_therapeute": ["Faites de petits pas pour faire confiance (test de réalité).", "Fixez des limites claires avec les personnes toxiques.", "Développez de la compassion pour l'enfant blessé en vous."],
        "action_pastorale": "Remplacez la suspicion systématique par la prière : 'Seigneur, donne-moi ton discernement'. Déposez les armes de la défensive à la Croix.",
        "verset": "Psaume 62:8 - « En tout temps, peuples, confiez-vous en lui, répandez vos cœurs en sa présence ! Dieu est notre refuge. »"
    },
    "SI": {
        "titre": "Isolement Social",
        "slogan": "« Je n'ai pas ma place ici »",
        "clinique_expert": "Sentiment de différence fondamentale ('Je suis un extraterrestre'). Exclusion du groupe, non par rejet actif, mais par manque d'appartenance ressentie. Le patient se vit comme fondamentalement inadapté au lien social.",
        "theologie_expert":"L'homme a été créé pour la communion. En Christ, la 'différence' n'est plus un motif d'exclusion mais une fonction dans le Corps (1 Corinthiens 12). La rédemption inclut la réintégration dans la famille de Dieu, brisant la malédiction de l'errance.",
        "origines": ["Humiliation ou rejet par les pairs (école).", "Famille différente de la communauté.", "Manque de compétences sociales encouragées."],
        "symptomes": ["Se sentir 'imposteur' ou 'inintéressant' en groupe.", "Évitement systématique des activités sociales.", "Caméléon social pour s'intégrer (perte de soi)."],
        "mecanisme_titre": "Le Cycle de l'Anxiété Sociale",
        "mecanisme_texte": "Anxiété -> Évitement -> Conséquences -> Renforcement de l'inadéquation -> Isolement accru.",
        "actions_therapeute": ["Exposez-vous progressivement aux situations évitées.", "Trouvez votre 'tribu' (intérêts communs).", "Entraînez-vous aux compétences sociales."],
        "action_pastorale": "Participez à la vie d'église non pour 'briller' mais pour 'être avec'. Vous êtes membre du Corps : l'œil ne peut dire à la main 'je n'ai pas besoin de toi'.",
        "verset": "Éphésiens 2:19 - « Vous n'êtes plus des étrangers, ni des gens du dehors; mais vous êtes concitoyens des saints. »"
    },
    "DS": {
        "titre": "Imperfection / Honte",
        "slogan": "« Je ne vaux rien »",
        "clinique_expert": "Sentiment d'être intrinsèquement défectueux (Badness). La honte est ici toxique : ce n'est pas 'j'ai fait une erreur' (culpabilité), mais 'je SUIS une erreur'. Cela entraîne une hypersensibilité à la critique et des stratégies de masque.",
        "theologie_expert": "C'est le cœur de l'Évangile : la Justification. Christ a pris notre honte à la croix. Nous sommes déclarés justes non par notre mérite ou par qui nous sommes, mais par l'imputation de sa justice, a cause de qui est Dieu. La valeur du patient ne dépend plus de son 'état', mais de son 'statut' en Christ.",
        "origines": ["Famille critique, humiliante ou punitive.", "Rejet ou manque d'amour par un parent.", "Comparaison défavorable avec la fratrie."],
        "symptomes": ["Cacher sa vraie personnalité (masque).", "Hypersensibilité à la critique.", "Attaquer les autres pour se revaloriser."],
        "mecanisme_titre": "Les 3 Copings de la Honte",
        "mecanisme_texte": "1. Capitulation (autodestruction). 2. Évitement (se cacher). 3. Contre-attaque (narcissisme/critique).",
        "actions_therapeute": ["Cessez de vous comparer aux autres.", "Dressez une liste de vos qualités réelles.", "Acceptez les compliments sans les minimiser."],
        "action_pastorale": "Quand la voix critique attaque, répondez à voix haute : 'Je suis imparfait, mais justifié, lavé et aimé en Christ'. Votre valeur a été fixée à la Croix.",
        "verset": "Sophonie 3:17 - « Il fera de toi le sujet de sa joie... il se réjouira à ton sujet avec des cris de joie. »"
    },
    "FA": {
        "titre": "Échec",
        "slogan": "« Je suis un raté »",
        "clinique_expert": "Croyance en l'incompétence relative aux pairs. Le patient s'identifie à ses échecs scolaires ou professionnels. Il y a souvent un évitement des défis pour ne pas confirmer cette croyance (prophétie auto-réalisatrice).",
        "theologie_expert": "L'idolâtrie de la réussite sociale est brisée. Dieu appelle souvent 'les choses faibles pour confondre les fortes'. Le succès selon le Royaume est la fidélité, pas le résultat visible. La dignité du travail est restaurée comme service, non comme prouesse.",
        "origines": ["Parents très critiques sur les résultats.", "Comparaison défavorable avec les autres enfants.", "Manque de limites ou de discipline dans l'enfance."],
        "symptomes": ["Procrastination par peur de l'échec.", "Minimiser ses propres réussites.", "Abandonner rapidement une tâche."],
        "mecanisme_titre": "La Pensée 'Tout ou Rien'",
        "mecanisme_texte": "Vision dichotomique : 'Si je ne suis pas le meilleur, je suis un échec total'. Cette norme irréaliste condamne à l'échec perçu.",
        "actions_therapeute": ["Reconnaissez la courbe d'apprentissage normale.", "Faites une liste de vos compétences réelles.", "Lancez un hobby sans enjeu de performance."],
        "action_pastorale": "Redéfinissez le succès : pour Dieu, c'est vivre pour sa gloire dans l'amour et l'obéissance. Entreprenez une action en acceptant qu'elle soit 'moyenne' aux yeux du monde, mais faite pour la gloire de Dieu.",
        "verset": "2 Corinthiens 12:9 - « Ma grâce te suffit, car ma puissance s'accomplit dans la faiblesse. »"
    },
    "DI": {
        "titre": "Dépendance / Incompétence",
        "slogan": "« Je n'y arrive pas tout seul »",
        "clinique_expert": "Croyance en l'incapacité à survivre seul. Le patient régresse dans une posture infantile, cherchant une 'figure parentale' pour assumer ses responsabilités. Manque de confiance dans son propre jugement et compétences (Self-Efficacy faible).",
        "theologie_expert": "Dieu nous a donné un esprit de force et de sagesse (2 Tim 1:7). La dépendance saine est verticale (envers Dieu), ce qui permet une autonomie horizontale (envers les hommes). L'Esprit Saint est le 'Paraclet' qui capacite le croyant à marcher.",
        "origines": ["Parents surprotecteurs ('je le fais pour toi').", "Parents qui ne laissaient pas prendre de décisions.", "Manque de conseils pratiques (négligence)."],
        "symptomes": ["Besoin constant d'être rassuré.", "Peur paralysante de prendre une mauvaise décision.", "Laisser les autres diriger sa vie."],
        "mecanisme_titre": "Surprotection vs Négligence",
        "mecanisme_texte": "Soit l'enfant a été étouffé (pas d'autonomie), soit il a été livré à lui-même trop tôt sans guidance (échec appris).",
        "actions_therapeute": ["Listez les tâches où vous dépendez des autres.", "Prenez des petites décisions seul et assumez le résultat.", "Célébrez chaque acte d'autonomie."],
        "action_pastorale": "Prenez une décision quotidienne seul(e) (repas, trajet) en vous confiant au Saint-Esprit qui habite en vous. Vous êtes équipé pour la vie. Dieu vous donne la force et les compétences dont vous avez besoin.",
        "verset": "Philippiens 4:13 - « Je puis tout par celui qui me fortifie. »"
    },
    "VU": {
        "titre": "Vulnérabilité au Danger",
        "slogan": "« Une catastrophe arrive »",
        "clinique_expert": "Anxiété catastrophique. Le monde est perçu comme un lieu de dangers imminents (maladie, ruine) qu'on ne peut ni prévoir ni contrôler. Hypervigilance constante du système nerveux.",
        "theologie_expert": "Le problème racine est le contrôle. L'anxiété est une tentative d'assumer la Souveraineté de Dieu. La paix vient non pas de la sécurité totale (impossible), mais de la confiance en la Providence divine qui tient les temps et les circonstances.",
        "origines": ["Parent anxieux ou phobique.", "Traumatisme ou maladie grave dans l'enfance.", "Environnement insécure."],
        "symptomes": ["Scénarios catastrophes.", "Vérifications compulsives.", "Rituels superstitieux."],
        "mecanisme_titre": "Distorsions Cognitives",
        "mecanisme_texte": "1. Catastrophisme (le pire va arriver).\n2. Surestimation du danger / Sous-estimation de ses capacités.\n3. Superstition (pensée magique).",
        "actions_therapeute": ["Analysez la probabilité réelle des catastrophes.", "Réduisez les comportements de vérification.", "Exposition progressive aux situations craintes."],
        "action_pastorale": "Faites une 'Diète de l'info' anxiogène. Tenez un carnet de Gratitude notant 3 événements par jour où vous avez été gardé par Dieu. Ancrez-vous dans la sécurité que Dieu vous donne dans le présent.",
        "verset": "Psaume 91:4 - « Il te couvrira de ses plumes, et tu trouveras un refuge sous ses ailes. »"
    },
    "EU": {
        "titre": "Fusion / Personnalité Atrophiée",
        "slogan": "« Je ne peux pas vivre sans toi »",
        "clinique_expert": "Symbiose émotionnelle. Le patient n'a pas achevé son processus d'individuation. Il vit par procuration, absorbant les émotions de l'autre. Sentiment de vide existentiel sans la figure d'attachement.",
        "theologie_expert": "Dieu a créé des individus distincts responsables de leurs propres âmes. La fusion est une forme d'idolâtrie relationnelle. Christ appelle à le suivre, ce qui nécessite parfois de 'quitter' (émotionnellement) père et mère pour devenir une personne entière.",
        "origines": ["Parent envahissant ne respectant pas les frontières.", "Culpabilisation quand l'enfant s'autonomise.", "Parent vivant à travers l'enfant."],
        "symptomes": ["Sentiment de vide quand on est seul.", "Imiter les émotions/avis de l'autre.", "Culpabilité intense à avoir une vie privée."],
        "mecanisme_titre": "Identité Non-Développée",
        "mecanisme_texte": "La personne ne sait pas qui elle est sans l'autre. Elle se définit par 'nous' plutôt que 'je'. Risque de relations toxiques.",
        "actions_therapeute": ["Listez vos préférences personnelles (goûts, avis) distincts de l'autre.", "Passez du temps seul pour découvrir qui vous êtes.", "Fixez des limites."],
        "action_pastorale": "Osez exprimer une opinion différente d'un proche sur un sujet mineur. C'est un acte spirituel d'affirmation de la créature unique que Dieu a faite en vous créant.",
        "verset": "Galates 1:10 - « Est-ce la faveur des hommes que je désire, ou celle de Dieu ? »"
    },
    "SB": {
        "titre": "Assujettissement",
        "slogan": "« Je dois faire ce que tu veux »",
        "clinique_expert": "Soumission forcée pour éviter la colère ou l'abandon. Le patient réprime ses besoins et accumule une colère latente (agressivité passive). Il ne se sent pas le 'droit' d'avoir des limites.",
        "theologie_expert": "La crainte de l'homme est un piège. Le chrétien est serviteur de Dieu, ce qui l'affranchit de l'esclavage des hommes. La vraie soumission est un choix libre d'amour (agapé), pas une contrainte de peur (phobos). Dire 'non' est parfois un acte spirituel.",
        "origines": ["Parent dominant, contrôlant ou punitif.", "Menaces si désaccord.", "Rôle de parentification."],
        "symptomes": ["Peur de dire non.", "Sentiment d'être piégé.", "Accumulation de colère (ressentiment)."],
        "mecanisme_titre": "Le Rôle de la Colère Refoulée",
        "mecanisme_texte": "La soumission crée une dette émotionnelle. La colère refoulée finit par exploser ou devenir des symptômes psychosomatiques.",
        "actions_therapeute": ["Entraînez-vous à dire 'non' sur des petites choses.", "Identifiez vos droits et besoins légitimes.", "Tolérer l'inconfort de ne pas plaire."],
        "action_pastorale": "Exercez-vous au 'Non bienveillant'. Refusez une demande cette semaine. Rappelez-vous que vous servez Dieu, pas l'humeur changeante des autres.",
        "verset": "Galates 5:1 - « C'est pour la liberté que Christ nous a affranchis. »"
    },
    "SS": {
        "titre": "Abnégation",
        "slogan": "« Je suis le sauveur »",
        "clinique_expert": "Le syndrome du Sauveur. Focalisation excessive sur les besoins d'autrui au détriment des siens, motivée par la culpabilité ou le besoin de valorisation narcissique ('Je suis utile donc je suis').",
        "theologie_expert": "Nous ne sommes pas le Messie. Vouloir sauver tout le monde est une limite que seul Dieu peut franchir. L'intendance (gérance) de son propre corps et de son âme est un devoir biblique. L'amour du prochain implique de s'aimer soi-même correctement.",
        "origines": ["Responsabilité excessive d'un proche dans l'enfance.", "Valorisée uniquement quand elle donnait.", "Tempérament hyper-empathique."],
        "symptomes": ["Ne pas savoir recevoir de l'aide.", "Épuisement (burnout).", "Attiré par les personnes à problèmes."],
        "mecanisme_titre": "Frontières (Boundaries)",
        "mecanisme_texte": "Difficulté à fixer des limites. Le sacrifice est motivé par la culpabilité, pas par l'amour libre. C'est une forme de codépendance.",
        "actions_therapeute": ["Équilibrez le donner et le recevoir.", "Demandez-vous : 'Je le fais par envie ou par culpabilité ?'.", "Acceptez que les autres gèrent leurs problèmes."],
        "action_pastorale": "Pratiquez le Sabbat : une demi-journée ou une journée sans 'servir', juste pour être aimé de Dieu sans rien faire. C'est un acte d'humilité : le monde tourne sans vous.",
        "verset": "Matthieu 22:39 - « Tu aimeras ton prochain comme toi-même. »"
    },
    "EI": {
        "titre": "Inhibition Émotionnelle",
        "slogan": "« Je ne dois pas ressentir »",
        "clinique_expert": "Sur-contrôle des affects. La spontanéité est jugée dangereuse ou honteuse. Le patient présente un 'faux-self' rationnel et froid pour se protéger de la vulnérabilité.",
        "theologie_expert": "Jésus a pleuré, a ressenti la colère et l'angoisse. Les émotions sont créées par Dieu comme des signaux. Les réprimer, c'est vivre dans le mensonge intérieur. La vérité (aletheia) implique l'authenticité émotionnelle devant Dieu (Psaumes de lamentation).",
        "origines": ["Émotions moquées ou punies.", "Famille puritaine ou très rationnelle.", "Peur de ressembler à un parent hystérique."],
        "symptomes": ["Paraître froid ou distant.", "Incapacité à pleurer ou montrer sa joie.", "Accent excessif sur la logique."],
        "mecanisme_titre": "La Roue des Émotions",
        "mecanisme_texte": "Inhibition de la colère, de la joie ou de la vulnérabilité. Tendance à rationaliser pour éviter de ressentir la douleur.",
        "actions_therapeute": ["Utilisez la 'Roue des émotions' pour nommer ce que vous ressentez.", "Tenez un journal émotionnel.", "Recherchez des expériences émotionnelles."],
        "action_pastorale": "Priez avec les Psaumes de lamentation. Osez dire 'Je suis triste' ou 'Je suis en colère' à Dieu. Les émotions ne sont pas des péchés, ce sont des informations.",
        "verset": "Psaume 62:9 - « Répandez votre cœur en sa présence ! Dieu est notre refuge. »"
    },
    "US": {
        "titre": "Exigences Élevées",
        "slogan": "« Ce n'est jamais assez bien »",
        "clinique_expert": "Perfectionnisme pathologique. La valeur personnelle est conditionnelle à la performance. Tyrannie du 'Je dois'. Incapacité à ressentir la satisfaction ou le repos.",
        "theologie_expert": "C'est une forme de légalisme : chercher à se justifier par les œuvres. L'Évangile est la fin de la performance pour le salut. Dieu a institué le Sabbat (repos) pour rappeler que le monde tourne sans nos efforts. La Grâce est l'acceptation de l'imperfection.",
        "origines": ["Amour conditionnel à la réussite.", "Parents perfectionnistes.", "Critique ou honte en cas d'échec."],
        "symptomes": ["Impossible de se détendre.", "Hyper-critique envers soi et les autres.", "Symptômes physiques de stress."],
        "mecanisme_titre": "Les 3 Types de Normes",
        "mecanisme_texte": "1. Compulsivité (ordre). 2. Orientation réussite (travail). 3. Orientation statut (image). C'est une course sans fin.",
        "actions_therapeute": ["Essayez de réduire vos exigences de 10%.", "Listez les avantages et inconvénients de votre pression.", "Forcez-vous à ralentir."],
        "action_pastorale": "Le défi de l'imperfection : laissez volontairement une tâche inachevée (ex: lit mal fait) et observez que Dieu vous aime toujours autant. La grâce suffit.",
        "verset": "Matthieu 11:28 - « Venez à moi, vous tous qui êtes fatigués et chargés, et je vous donnerai du repos. »"
    },
    "ET": {
        "titre": "Droits Personnels / Grandeur",
        "slogan": "« Les règles ne s'appliquent pas à moi »",
        "clinique_expert": "Narcissisme et sentiment de privilège. Le patient refuse les limites communes, manque d'empathie et tolère mal la frustration. C'est souvent une compensation d'un sentiment d'infériorité caché.",
        "theologie_expert": "L'orgueil précède la chute. Le Royaume de Dieu est un 'monde à l'envers' où le plus grand est le serviteur. Reconnaître sa dépendance totale à la grâce de Dieu est le seul remède à l'inflation de l'ego. L'autre n'est pas un outil, mais un porteur de l'image de Dieu.",
        "origines": ["Enfant gâté, sans limites.", "Parents n'ayant pas imposé de conséquences.", "Compensation d'un sentiment de manque."],
        "symptomes": ["Colère si on ne l'obéit pas immédiatement.", "Manque d'empathie.", "Compétitivité excessive."],
        "mecanisme_titre": "Les 3 Types de Droits",
        "mecanisme_texte": "1. Narcissisme pur (je suis spécial). 2. Dépendance (les autres doivent me servir). 3. Impulsivité (je veux tout, tout de suite).",
        "actions_therapeute": ["Mettez-vous à la place des autres (empathie cognitive).", "Demandez un feedback honnête.", "Respectez les règles communes."],
        "action_pastorale": "Pratiquez le service anonyme. Faites une bonne action (vaisselle, don) sans le dire et sans attendre de merci. Écoutez les autres sans ramener la conversation à vous.",
        "verset": "Philippiens 2:3 - « Regardez les autres comme étant au-dessus de vous-mêmes. »"
    },
    "IS": {
        "titre": "Contrôle de soi insuffisant",
        "slogan": "« C'est trop difficile, je m'en fiche »",
        "clinique_expert": "Impulsivité et intolérance à la frustration. Le principe de plaisir domine le principe de réalité. Difficulté à différer la gratification pour un but à long terme.",
        "theologie_expert": "La maîtrise de soi est un fruit de l'Esprit (Galates 5). Ce n'est pas une simple volonté humaine, mais une discipline spirituelle. C'est apprendre à dire 'non' à la chair pour dire 'oui' à la vie. La sagesse biblique valorise la construction patiente.",
        "origines": ["Manque de discipline parentale.", "Négligence ou stress chronique.", "Enfant jamais forcé à tolérer la frustration."],
        "symptomes": ["Procrastination chronique.", "Addictions.", "Évitement systématique de l'inconfort."],
        "mecanisme_titre": "La Stratégie SNAP",
        "mecanisme_texte": "Stop (arrêtez-vous dans le comportement impulsif), Noter (Remarquer ce qui se passe autour de vous et en vous, ce qui déclanche votre réaction), Aligner (Aligner vos actions avec vos objectif et valeurs), Processer (Agissez conformement à vos valeurs et objectifs). Outil pour briser l'impulsion.",
        "actions_therapeute": ["Utilisez la méthode SNAP.", "Fixez des micro-objectifs réalisables.", "Enlevez les distractions."],
        "action_pastorale": "La méthode des 10 minutes : Quand vous voulez abandonner une tâche ou céder à une impulsion, tenez 10 min de plus en priant. C'est un muscle spirituel à exercer.",
        "verset": "Proverbes 25:28 - « Comme une ville forcée et sans murailles, ainsi est l'homme qui n'est pas maître de lui-même. »"
    },
    "AS": {
        "titre": "Recherche d'approbation",
        "slogan": "« Ma valeur dépend de ton regard »",
        "clinique_expert": "Le 'Caméléon'. L'estime de soi est externalisée : elle dépend entièrement du regard de l'autre. Le patient perd son authenticité pour s'adapter aux attentes supposées de l'entourage.",
        "theologie_expert": "La crainte de l'homme est un piège. C'est de l'idolâtrie de l'approbation. Le chrétien vit 'Coram Deo' (devant la face de Dieu). Seule l'approbation du Père ('Tu es mon fils bien-aimé') peut saturer ce besoin et libérer de la tyrannie du regard d'autrui.",
        "origines": ["Amour conditionné à la 'bonne conduite'.", "Parents soucieux des apparences.", "Manque d'attention."],
        "symptomes": ["Changer de personnalité.", "Importance excessive du statut.", "Peur panique de déplaire."],
        "mecanisme_titre": "L'Adaptation Excessive",
        "mecanisme_texte": "Le patient pense : 'Si je suis moi-même, on ne m'aimera pas'. Il développe un 'Faux-Self' pour être validé.",
        "actions_therapeute": ["Demandez-vous : 'Qu'est-ce que JE veux ?'.", "Entraînez-vous à exprimer vos préférences.", "Passez du temps seul."],
        "action_pastorale": "Faites le bien en secret (Matthieu 6). Acceptez un compliment par un simple 'Merci' sans vous en nourrir excessivement ni le rejeter. L'approbation la plus importante est celle de Dieu seul.",
        "verset": "1 Thessaloniciens 2:4 - « Nous parlons, non pour plaire aux hommes, mais pour plaire à Dieu. »"
    },
    "NP": {
        "titre": "Négativité / Pessimisme",
        "slogan": "« Ça va mal finir »",
        "clinique_expert": "Biais cognitif de focalisation sur le négatif. Attente anxieuse que 'tout va s'effondrer'. Le positif est minimisé ou considéré comme suspect. Souvent lié à une anxiété chronique.",
        "theologie_expert": "Bien que le mal soit réel, la résignation est un déni de la bonté de Dieu et de l'Espérance. La 'joie' biblique est un combat de la foi, une discipline de l'attention (Phil 4:8) pour reconnaître la grâce commune et la providence au milieu des épreuves.",
        "origines": ["Parents pessimistes ou inquiets.", "Enfance marquée par des difficultés.", "Découragement de l'autonomie."],
        "symptomes": ["Filtre négatif.", "Incapacité à se réjouir.", "Plaintes chroniques."],
        "mecanisme_titre": "Les Distorsions Cognitives",
        "mecanisme_texte": "1. Filtre négatif. 2. Généralisation excessive ('toujours'). 3. Catastrophisme. C'est une protection : 'Si je m'attends au pire, je ne serai pas déçu'.",
        "actions_therapeute": ["Examinez les preuves.", "Tenez un journal de gratitude.", "Considérez les exceptions."],
        "action_pastorale": "Contre la rumination, trouvez un aspect positif pour chaque pensée négative. Louez Dieu pour une petite chose précise chaque matin.",
        "verset": "Lamentations 3:21 - « Voici ce que je veux repasser en mon cœur, ce qui me donnera de l'espérance. »"
    },
    "PU": {
        "titre": "Punition",
        "slogan": "« Les erreurs doivent être punies »",
        "clinique_expert": "Intransigeance et dureté. Croyance que l'erreur mérite châtiment. Difficulté à pardonner (à soi et aux autres). Tendance au jugement moralisateur.",
        "theologie_expert": "C'est une incompréhension de la Croix. Christ a pris la punition. Il n'y a plus de condamnation (Rom 8:1). Maintenir une attitude punitive, c'est nier la suffisance du sacrifice de Jésus. Nous sommes appelés à être des canaux de la miséricorde que nous avons reçue.",
        "origines": ["Punitions sévères dans l'enfance.", "Parents impitoyables.", "Manque de droit à l'erreur."],
        "symptomes": ["Rancune tenace.", "Autopunition.", "Jugement sévère des autres."],
        "mecanisme_titre": "Le Cycle de la Rancune",
        "mecanisme_texte": "Standards rigides -> Erreur inévitable -> Colère/Jugement -> Punition. Croyance que la punition 'corrige' le comportement.",
        "actions_therapeute": ["Pratiquez l'auto-compassion.", "Considérez les circonstances atténuantes.", "Pardonnez-vous une erreur passée."],
        "action_pastorale": "Si Jésus a payé, ne cherchez pas à payer encore. Parlez-vous avec la douceur que le Christ utilise pour vous parler.",
        "verset": "Romains 8:1 - « Il n'y a donc maintenant aucune condamnation pour ceux qui sont en Jésus-Christ. »"
    }
}

YOUNG_DOMAINS_INFO = {
    "Domaine I : Séparation et Rejet": {"codes": ["ED", "AB", "MA", "SI", "DS"], "besoin": "Besoin de sécurité, de stabilité, d'affection et d'appartenance."},
    "Domaine II : Manque d'Autonomie": {"codes": ["DI", "VU", "EU", "FA"], "besoin": "Besoin de compétence, d'identité propre et de confiance en soi."},
    "Domaine III : Limites Déficientes": {"codes": ["ET", "IS"], "besoin": "Besoin de limites réalistes, de respect des autres et d'autodiscipline."},
    "Domaine IV : Orientation vers les Autres": {"codes": ["SB", "SS", "AS"], "besoin": "Besoin de liberté d'expression et d'affirmation de ses besoins."},
    "Domaine V : Hypervigilance et Inhibition": {"codes": ["NP", "EI", "US", "PU"], "besoin": "Besoin de spontanéité, de plaisir et de lâcher-prise."}
}
# --- LES 232 QUESTIONS VALIDÉES (OFFICIELLES) ---
YSQ_QUESTIONS = {
    "ED : Carence affective": {
        1: "Les autres n'ont pas satisfait mes besoins affectifs.",
        2: "Je n'ai pas reçu suffisamment d'amour et d'attention.",
        3: "Dans l'ensemble, je n'ai eu personne sur qui compter pour recevoir des conseils ou du soutien affectif.",
        4: "La plupart du temps, je n'ai eu personne pour m'aider à grandir, pour partager son univers intérieur avec moi, ou qui se soucie profondément de tout ce qui m'arrive.",
        5: "Pour la plus grande partie de ma vie, je n'ai eu personne qui veuille être près de moi et passer beaucoup de temps avec moi.",
        6: "En général, les autres n'ont pas été présents pour me prendre dans leurs bras, pour me donner de la chaleur et de l'affection.",
        7: "Pour la plus grande partie de ma vie, je n'ai jamais eu le sentiment que je représentais quelqu'un d'important pour quelqu'un d'autre.",
        8: "En grande partie, je n'ai eu personne qui m'écoute réellement, me comprenne et soit sensible à mes besoins et mes sentiments véritables.",
        9: "Je n'ai pas eu une personne forte pour me donner de bons conseils ou pour me guider lorsque je ne savais pas quoi faire."
    },
    "AB : Abandon / Instabilité": {
        10: "Je suis préoccupé(e) par le fait que les gens que j'aime vont mourir bientôt même s'il y a peu de raisons médicales à ma préoccupation.",
        11: "Je m'accroche aux gens dont je suis proche par peur qu'ils ne me quittent.",
        12: "Je crains que les gens dont je me sens proche ne me quittent ou ne m'abandonnent.",
        13: "J'ai le sentiment de manquer d'une base stable qui me soutienne affectivement.",
        14: "Je n'ai pas l'impression que les relations importantes dureront: je m'attends à ce qu'elles finissent.",
        15: "Je me sens 'accro' aux partenaires qui ne peuvent pas s'engager avec moi de façon stable.",
        16: "Je finirai seul(e).",
        17: "Quand je sens que quelqu'un à qui je tiens s'éloigne de moi, je deviens désespéré(e).",
        18: "Quelquefois j'ai tellement peur que les gens m'abandonnent que je les repousse.",
        19: "Je deviens bouleversé(e) quand quelqu'un me laisse seul(e) même pour une courte période.",
        20: "Je ne peux compter sur la présence régulière de ceux qui me soutiennent.",
        21: "Je ne peux me permettre d'être vraiment proche des autres car je ne peux être sûr(e) qu'ils seront toujours là.",
        22: "Il me semble que les personnes importantes dans ma vie sont toujours en train de venir et de repartir.",
        23: "J'ai très peur que les personnes que j'aime ne trouvent quelqu'un d'autre qu'elles préfèrent et qu'elles m'abandonnent.",
        24: "Les gens qui me sont proches ont toujours été très imprévisibles.",
        25: "J'ai tellement besoin des autres que j'ai peur de les perdre.",
        26: "Je ne peux être moi-même ou exprimer ce que je ressens véritablement, sinon les autres vont me quitter."
    },
    "MA : Méfiance / Abus": {
        27: "J'ai l'impression que les autres vont profiter de moi.",
        28: "J'ai souvent l'impression que je dois me protéger des autres.",
        29: "J'ai l'impression que je dois être sur mes gardes en présence des autres sinon ils me blesseront intentionnellement.",
        30: "Si une personne est aimable avec moi, je suppose qu'elle cherche à obtenir quelque chose.",
        31: "Ce n'est qu'une question de temps avant que quelqu'un me trahisse.",
        32: "La plupart des gens pensent uniquement à eux.",
        33: "J'ai la plus grande difficulté à faire confiance aux autres.",
        34: "Je suis très méfiant quant aux motivations des autres.",
        35: "Les autres sont rarement honnêtes, ils ne sont pas en général ce qu'ils paraissent.",
        36: "Je m'interroge habituellement sur les véritables intentions des autres.",
        37: "Si je pense que quelqu'un cherche à me blesser, je cherche à le blesser en premier.",
        38: "Les autres habituellement doivent faire leurs preuves avant que je leur accorde ma confiance.",
        39: "Je teste les autres pour voir s'ils me disent la vérité ou s'ils sont bien intentionnés.",
        40: "Je souscris à la croyance: 'Contrôle ou tu seras contrôlé(e)'.",
        41: "Je me mets en colère quand je pense aux façons dont j'ai été maltraité(e) par les autres tout au long de ma vie.",
        42: "Tout au long de ma vie, mes proches ont profité de moi et m'ont utilisé(e) à leurs propres desseins.",
        43: "J'ai été physiquement, émotionnellement, ou sexuellement abusé(e) par des personnes importantes de ma vie."
    },
    "SI : Isolement social": {
        44: "Je ne suis pas adapté(e).",
        45: "Je suis fondamentalement différent(e) des autres.",
        46: "Je suis à part; je suis un(e) solitaire.",
        47: "Je me sens étranger(ère) aux autres.",
        48: "Je me sens isolé(e) et seul(e).",
        49: "Je me sens toujours à l'extérieur des groupes.",
        50: "Personne ne me comprend vraiment.",
        51: "Ma famille a toujours été différente des autres familles.",
        52: "J'ai parfois le sentiment d'être un étranger.",
        53: "Si je disparaissais demain, personne ne le remarquerait."
    },
    "DS : Imperfection / Honte": {
        54: "Aucun homme ou femme que je désire ne pourrait m'aimer une fois qu'il(elle) aurait vu mes défauts.",
        55: "Aucune personne que je désire ne pourrait rester proche de moi si elle savait qui je suis réellement.",
        56: "Je suis fondamentalement imparfait(e) et marqué(e) par un défaut.",
        57: "Même si je me donne le plus grand mal, je sens qu'il ne me sera pas possible d'obtenir le respect d'un homme ou d'une femme important(e) et de sentir que j'ai de la valeur.",
        58: "Je ne mérite pas l'amour, l'attention et le respect des autres.",
        59: "J'ai le sentiment d'être quelqu'un que l'on ne peut pas aimer.",
        60: "Je suis trop fondamentalement inacceptable pour me révéler aux autres.",
        61: "Je ne pourrais pas affronter les gens, s'ils découvraient mes défauts fondamentaux.",
        62: "Lorsque les gens m'apprécient, j'ai l'impression de les duper.",
        63: "Je suis souvent attiré(e) par les gens qui sont très critiques envers moi ou qui me rejettent.",
        64: "J'ai des secrets que je ne veux pas que mes proches découvrent.",
        65: "C'est de ma faute si mes parents n'ont pas pu m'aimer suffisamment.",
        66: "Je ne laisse pas les gens connaître ce que je suis réellement.",
        67: "Une de mes plus grandes peurs est que mes défauts deviennent publics.",
        68: "Je ne puis comprendre comment qui que ce soit pourrait m'aimer."
    },
    "FA : Échec": {
        69: "Presque rien de ce que je fais au travail (ou à l'école) n'est aussi bon que ce que font les autres.",
        70: "Je suis incompétent(e) quand il s'agit de réussir.",
        71: "La plupart des gens sont plus doués que moi en ce qui concerne le travail (ou l'école) et la réussite.",
        72: "Je suis un(e) raté(e).",
        73: "Je n'ai pas autant de talent que les autres au travail (ou à l'école).",
        74: "Je ne suis pas aussi intelligent(e) que la plupart des gens quand il s'agit du travail (ou de l'école).",
        75: "Je suis humilié(e) par mes échecs ou mes insuffisances dans le monde du travail(ou de l'école).",
        76: "Je suis souvent mal à l'aise avec les autres, car je ne les vaux pas en termes de réussites.",
        77: "Je compare souvent mes réalisations à celles des autres et je trouve qu'ils réussissent beaucoup mieux."
    },
    "DI : Dépendance / Incompétence": {
        78: "Je ne me sens pas capable de me débrouiller par moi-même dans la vie de tous les jours.",
        79: "J'ai besoin des autres pour m'aider à m'en sortir.",
        80: "Je n'ai pas le sentiment que je puisse bien m'adapter par moi-même.",
        81: "Je crois que les autres peuvent prendre soin de moi mieux que je ne le peux moi-même.",
        82: "J'ai des difficultés à prendre en charge de nouvelles tâches en dehors du travail à moins que quelqu'un ne me guide.",
        83: "Je me considère comme une personne dépendante en ce qui concerne la vie de tous les jours.",
        84: "Je bousille tout ce que j'entreprends, même à l'extérieur du travail (ou de l'école).",
        85: "Je suis stupide dans la plupart des domaines de la vie.",
        86: "Si je me fie à mon jugement dans la vie de tous les jours, je vais prendre la mauvaise décision.",
        87: "Je manque de bon sens.",
        88: "On ne peut se fier à mon jugement dans les situations quotidiennes.",
        89: "Je n'ai pas confiance dans ma capacité à résoudre les problèmes qui se posent tous les jours.",
        90: "Je pense avoir besoin de quelqu'un sur qui je puisse compter pour me donner des conseils sur les questions pratiques.",
        91: "Je me sens plus un(e) enfant qu'un(e) adulte quand il s'agit de prendre en main les responsabilités quotidiennes.",
        92: "Je me sens dépassé(e) par les responsabilités de tous les jours."
    },
    "VU : Vulnérabilité": {
        93: "Il ne me semble pas possible d'échapper au sentiment que quelque chose de mauvais va bientôt se passer.",
        94: "J'ai l'impression qu'un désastre naturel, criminel, financier ou médical pourrait frapper à tout moment.",
        95: "J'ai peur de devenir un(e) sans domicile fixe ou un(e) mendiant(e).",
        96: "J'ai peur d'être attaqué(e).",
        97: "Je prends de grandes précautions pour éviter de tomber malade ou d'être blessé(e).",
        98: "J'ai peur d'avoir une maladie grave, même si rien de sérieux n'a été diagnostiqué par un médecin.",
        99: "Je suis quelqu'un de peureux(se).",
        100: "Je me soucie beaucoup de ce qui va mal dans le monde: le crime, la pollution, etc.",
        101: "J'ai souvent le sentiment que je pourrais devenir fou(folle).",
        102: "J'ai souvent l'impression que je vais avoir une crise d'angoisse.",
        103: "J'ai souvent peur d'avoir une crise cardiaque ou un cancer, même s'il y a peu de raisons médicales de s'en soucier.",
        104: "Je pense que le monde est un endroit dangereux."
    },
    "EU : Fusion / Personnalité atrophiée": {
        105: "Je n'ai pas pu me séparer de ma mère ou de mon père comme semblent le faire les gens de mon âge.",
        106: "Mes parents et moi avons tendance à être sur-impliqués dans nos vies et nos problèmes réciproques.",
        107: "Il est très difficile, pour mes parents et moi-même, de garder secrets, chacun pour soi, certains détails intimes, sans nous sentir trahis ou coupables.",
        108: "Mes parents et moi devons nous parler presque tous les jours, sinon l'un de nous se sent coupable, blessé(e), déçu(e), ou seul(e).",
        109: "J'ai souvent l'impression de ne pas avoir une identité distincte de celle de mes parents ou de mon partenaire.",
        110: "J'ai souvent l'impression que mes parents vivent à travers moi -- je n'ai pas une vie qui me soit propre.",
        111: "Il m'est vraiment ardu de maintenir une distance vis-à-vis des gens dont je suis intime; il m'est difficile de me sentir une personne séparée.",
        112: "Je suis tellement lié(e) à mon partenaire ou à mes parents que je ne sais pas vraiment qui je suis ou ce que je veux.",
        113: "J'éprouve des difficultés à distinguer mon point de vue ou mon opinion de ceux de mes parents ou de mon partenaire.",
        114: "J'ai souvent l'impression de ne pas avoir d'intimité par rapport à mes parents ou mon partenaire.",
        115: "Je sens que mes parents seraient vraiment peinés si je vivais seul(e), loin d'eux."
    },
    "SB : Assujettissement": {
        116: "Je laisse les autres faire ce qu'ils veulent, car j'ai peur des conséquences.",
        117: "Je pense que si je fais ce que je veux, je cours après les problèmes.",
        118: "Je sens que je n'ai pas d'autre choix que de me soumettre aux souhaits des autres, sinon ils exerceront des représailles ou me rejetteront d'une façon ou d'une autre.",
        119: "Dans mes relations, je laisse l'autre avoir le dessus sur moi.",
        120: "Je laisse toujours les autres choisir à ma place, si bien que je ne sais pas vraiment ce que je veux moi-même.",
        121: "J'ai le sentiment que les décisions importantes de ma vie n'étaient pas vraiment les miennes.",
        122: "Je me soucie beaucoup de plaire aux autres, pour qu'ils ne me rejettent pas.",
        123: "J'ai beaucoup de difficultés à exiger que mes droits soient respectés et que mes sentiments soient pris en compte.",
        124: "Plutôt que manifester ouvertement ma colère, je me venge dans des petites choses.",
        125: "Je vais tolérer beaucoup plus de choses que la plupart des gens afin d'éviter une confrontation."
    },
    "SS : Abnégation": {
        126: "Je fais passer les besoins des autres avant les miens, sinon je me sens coupable.",
        127: "Je me sens coupable si je laisse tomber les autres ou si je les déçois.",
        128: "Je donne davantage aux autres que je ne reçois en retour.",
        129: "Peu importe combien je donne, ce n'est jamais assez.",
        130: "Je suis celui (celle) qui finit généralement par prendre soin des gens dont je suis proche.",
        131: "Il n'y a presque rien que je ne puisse supporter lorsque j'aime quelqu'un.",
        132: "Je suis quelqu'un de bon car je pense aux autres plus qu'à moi-même.",
        133: "Au travail, je suis habituellement celui (celle) qui est volontaire pour faire des heures ou des tâches supplémentaires.",
        134: "Même si je suis très occupé(e), je trouve toujours du temps pour les autres.",
        135: "Je peux m'en sortir avec vraiment très peu car mes besoins sont minimes.",
        136: "Je ne suis heureux(se) que si les gens qui m'entourent le sont aussi.",
        137: "Je suis tellement occupé(e) à me dévouer pour les gens qui m'importent que j'ai très peu de temps pour moi.",
        138: "J'ai toujours été celui (celle) qui écoute les problèmes des autres.",
        139: "Je me sens plus à l'aise pour donner un cadeau que pour en recevoir un.",
        140: "On me voit comme quelqu'un qui en fait trop pour les autres et pas assez pour lui-même.",
        141: "Si je fais ce que je veux, je me sens vraiment mal à l'aise.",
        142: "Il m'est vraiment difficile de demander aux autres de se soucier de mes besoins."
    },
    "EI : Inhibition émotionnelle": {
        143: "J'ai peur de perdre le contrôle de mes actes.",
        144: "J'ai peur de faire du mal à quelqu'un, physiquement ou affectivement, si je perds tout contrôle sur ma colère.",
        145: "Je sens que je doit contrôler mes émotions et mes impulsions, sinon quelque chose de déplorable risque de se produire.",
        146: "J'en arrive à accumuler en moi beaucoup de colère et de ressentiment que je n'exprime pas.",
        147: "Je suis trop gêné(e) pour exprimer des sentiments positifs aux autres (par ex. de l'affection, de l'intérêt).",
        148: "Je trouve embarrassant d'exprimer mes sentiments aux autres.",
        149: "Il m'est difficile d'être chaleureux(se) et spontané(e).",
        150: "Je me contrôle tellement bien que les autres croient que je n'ai pas d'émotions.",
        151: "Les gens me trouvent coincé(e) sur le plan émotionnel."
    },
    "US : Exigences élevées": {
        152: "Je doit être le(la) meilleur(e) dans presque tout ce que je fais, je ne peux pas accepter d'être le(la) deuxième.",
        153: "Je m'efforce de tout maintenir parfaitement en ordre.",
        154: "Je doit apparaître sous mon meilleur aspect la plupart du temps.",
        155: "Je m'efforce de faire de mon mieux; je ne peux pas me contenter d'être « assez bien ».",
        156: "J'ai tant de choses à faire qu'il ne me reste presque pas de temps pour me détendre vraiment.",
        157: "Presque rien de ce que je fais n'est assez bien, je pourrais toujours faire mieux.",
        158: "Je doit faire face à toutes mes responsabilités.",
        159: "Je ressens une pression constante qui me force à faire les choses et à les réussir.",
        160: "Mes relations souffrent de ce que je m'impose trop de choses.",
        161: "Je m'impose une telle pression pour bien faire que ma santé en souffre.",
        162: "Souvent je sacrifie plaisir et bonheur pour atteindre mes idéaux.",
        163: "Lorsque je fais une erreur, je mérite une critique sévère.",
        164: "Je ne peux pas accepter de me dégager aisément d'une situation difficile ou de présenter des excuses pour mes erreurs.",
        165: "Je suis quelqu'un de très compétitif.",
        166: "J'attache beaucoup d'importance à l'argent ou au statut social.",
        167: "Je tiens toujours à être le(la) meilleur(e) dans ce que j'accomplis."
    },
    "ET : Droits personnels / Grandeur": {
        168: "Lorsque j'attends quelque chose des autres, j'ai beaucoup de difficulté à accepter un refus.",
        169: "Je me mets souvent en colère ou je m'irrite quand je ne peux pas avoir ce que je veux.",
        170: "Je suis quelqu'un de spécial et je ne devrais pas avoir à accepter les restrictions auxquelles les autres doivent se soumettre.",
        171: "Je déteste être contraint(e) ou qu'on m'empêche de faire ce que je veux.",
        172: "Je crois que je n'ai pas à suivre les règles et les conventions comme les autres.",
        173: "J'ai le sentiment que j'ai beaucoup plus à offrir que les autres.",
        174: "Habituellement, je fais passer mes besoins avant ceux des autres.",
        175: "Je remarque souvent que l'importance accordée à mes priorités personnelles ne me laisse plus de temps pour les amis ou la famille.",
        176: "Les autres me tells souvent que je contrôle trop la façon dont les choses sont faites.",
        177: "Je suis très irrité(e) lorsque les autres ne font pas ce que je leur demande.",
        178: "Je ne supporte pas que les autres me disent ce que je dois faire."
    },
    "IS : Contrôle de soi insuffisant": {
        179: "J'ai beaucoup de difficultés à m'arrêter de boire, de fumer, de trop manger ou à cesser d'autres comportements problématiques.",
        180: "Il ne me semble pas possible de me discipliner pour terminer des tâches routinières ou ennuyeuses.",
        181: "Souvent, je me laisse aller à mes impulsions ou à exprimer des émotions qui me créent des difficultés ou blessent les autres.",
        182: "Si je ne peux pas atteindre un but, je suis facilement frustré(e) et j'abandonne.",
        183: "Il m'est très difficile de sacrifier une gratification immédiate pour mener à bien un projet à long terme.",
        184: "Quand je me mets en colère, il est fréquent que je ne puisse pas du tout me contrôler.",
        185: "J'ai tendance à abuser de certaines choses, même si je sais que c'est mauvais pour moi.",
        186: "Je m'ennuie très facilement.",
        187: "Quand les tâches deviennent difficiles, je ne peux souvent pas persévérer pour les terminer.",
        188: "Je ne peux pas me concentrer longtemps sur quoi que ce soit.",
        189: "Je ne peux pas me forcer à faire des choses qui ne me plaisent pas même quand je sais que c'est pour mon bien.",
        190: "Je me mets en colère à la moindre offense.",
        191: "J'ai rarement été capable de tenir mes engagements.",
        192: "Je ne peux presque jamais m'empêcher de montrer aux gens mes véritables sentiments, même si cela doit me coûter cher.",
        193: "J'agis souvent impulsivement et je le regrette plus tard."
    },
    "AS : Recherche d'approbation": {
        194: "Il m'est important d'être aimé(e) par presque tous ceux que je connais.",
        195: "Je modifie mon comportement en fonction des gens avec qui je me trouve, afin qu'ils puissent m'aimer davantage.",
        196: "Je fais tout mon possible pour m'adapter.",
        197: "Mon estime personnelle est principalement fondée sur la façon dont les autres me voient.",
        198: "Posséder de l'argent et connaître des gens importants sont des choses qui me donnent de la valeur.",
        199: "Je consacre beaucoup de temps à soigner mon apparence afin de gagner l'estime des autres.",
        200: "Mes réussites ont plus de valeur à mes yeux si les autres les remarquent.",
        201: "Je suis à ce point soucieux(e) de m'adapter aux autres qu'il m'arrive de ne plus savoir qui je suis.",
        202: "Je trouve difficile de me fixer des objectifs propres, sans prendre en compte ce que les autres vont penser de mes choix.",
        203: "Lorsque je considère les décisions que j'ai prises dans ma vie, je réalise que la plupart d'entre elles reposaient sur le désir d'obtenir l'approbation des autres.",
        204: "Même si je n'aime pas une personne, je tiens malgré tout à ce qu'elle m'aime.",
        205: "À moins d'obtenir beaucoup d'attention des autres, je me sens peu important(e).",
        206: "Si j'interviens lors d'une réunion ou si on me présente dans un groupe, je cherche à être reconnu(e) et admiré(e).",
        207: "Lorsqu'on me fait beaucoup de compliments et d'éloges, je me sens une personne de valeur."
    },
    "NP : Négativité / Pessimisme": {
        208: "Même lorsque tout va bien, j'ai l'impression que ce ne sera que temporaire.",
        209: "S'il se produit quelque chose de bien, j'ai peur qu'il n'arrive ensuite quelque chose de mauvais.",
        210: "On n'est jamais assez prudent; il peut toujours se produire quelque chose de mauvais.",
        211: "Même si je travaille beaucoup, j'ai peur de me retrouver un jour sans le sou.",
        212: "J'ai peur qu'une mauvaise décision ne puisse conduire à un désastre.",
        213: "Je me tourmente souvent pour des décisions mineures, car les conséquences d'une erreur m'apparaissent tellement graves.",
        214: "Je préfère considérer que les choses ne vont pas aller comme je le veux, car ainsi je ne serai pas déçu(e) si ça se passe mal.",
        215: "Je pense plutôt aux aspects négatifs de la vie et des évènements qu'aux côtés positifs.",
        216: "J'ai tendance à être pessimiste.",
        217: "Mes proches considèrent que je me fais trop de soucis.",
        218: "Si les gens s'enthousiasment trop, je me sens mal et j'éprouve le besoin de les prévenir de ce qui pourrait mal se passer."
    },
    "PU : Punition": {
        219: "Si je fais une erreur, je mérite d'être puni(e).",
        220: "Si je ne donne pas le meilleur de moi-même, je dois m'attendre à échouer.",
        221: "Je n'ai aucune excuse si je fais une erreur.",
        222: "Les gens qui ne font pas leur part de travail à fond devraient être punis d'une manière ou d'une autre.",
        223: "La plupart du temps, je n'accepte pas les excuses des autres: ils n'ont pas pris leurs responsabilités et ils en paient les conséquences.",
        224: "Si je ne fais pas mon travail, je devrai en subir les conséquences.",
        225: "Je pense souvent à mes erreurs passées et je me mets en colère contre moi-même.",
        226: "Lorsque les gens font quelque chose de mal, j’ai du mal à appliquer la phrase: «pardonnez et oubliez».",
        227: "Je garde de la rancune envers les gens, même s’ils se sont excusés.",
        228: "Je me sens énervé(e) à l'idée que quelqu'un s'est tiré trop facilement d'un mauvais pas.",
        229: "Je me mets en colère lorsque des gens se trouvent des excuses, ou lorsqu’ils accusent les autres pour des problèmes dont ils sont eux-mêmes responsables.",
        230: "Peu importe les raisons, quand je fais une erreur, je devrais en payer le prix.",
        231: "Je m’en veux énormément lorsque je bousille quelque chose.",
        232: "Je suis quelqu’un qui mérite d’être puni quand il fait quelque chose de mal."
    }
}
# --- 4. INTERFACE ---
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/3050/3050525.png", width=80)
st.sidebar.title("Navigation")
mode = st.sidebar.radio("Aller vers :", ["Espace Patient", "Espace Thérapeute"])

if mode == "Espace Patient":
    st.header("Bienvenue dans votre questionnaire des Schémas (YSQ-L3)")
    st.markdown("---")
    st.info("""
    ### 💡 Guide pour remplir ce questionnaire
    Ce questionnaire est un outil précieux pour mieux comprendre votre fonctionnement émotionnel. Il ne s'agit pas d'un examen, mais d'une "photographie" de vos ressentis.
    
    **Comment répondre ?**
    1.  **Soyez spontané(e) :** Ne réfléchissez pas trop longtemps. Votre première impression est souvent la plus juste.
    2.  **Visez la globalité :** Répondez en fonction de ce que vous ressentez **la plupart du temps** dans votre vie, et pas seulement aujourd'hui.
    
    **L'échelle de notation :**
    * **1** : Ceci est **complètement faux** pour moi.
    * **2** : C'est **faux dans l'ensemble**, cela ne me ressemble pas vraiment.
    * **3** : C'est **plutôt vrai** que faux.
    * **4** : C'est **modérément vrai**, cela me correspond assez souvent.
    * **5** : C'est **vrai dans l'ensemble**, cela me décrit bien.
    * **6** : Ceci me décrit **parfaitement**, c'est tout à fait moi.
    
    *Vos réponses sont strictement confidentielles et seront analysées uniquement par votre thérapeute.*
    """)
    
    with st.form("form_patient", clear_on_submit=False):
        c1, c2 = st.columns(2)
        nom = c1.text_input("Votre Nom et Prénom *")
        email = c2.text_input("Votre Email *")
        reponses = {}
        st.divider()
        
        for i, (domaine, q_dict) in enumerate(YSQ_QUESTIONS.items()):
            with st.container():
                st.markdown(f"#### 📝 Série {i+1}")
                for q_num, q_text in q_dict.items():
                    st.write(f"**{q_num}.** {q_text}")
                    reponses[f"Q{q_num}"] = st.pills(f"Choix Q{q_num}", options=[1, 2, 3, 4, 5, 6], selection_mode="single", label_visibility="collapsed", key=f"q_{q_num}")
                    st.caption("")
            st.divider()
        
        if st.form_submit_button("Envoyer mes résultats", type="primary"):
            missing = [k for k, v in reponses.items() if v is None]
            if not nom or not email: st.error("⚠️ Merci de remplir votre nom et email.")
            elif missing: st.warning(f"⚠️ Il manque {len(missing)} réponse(s).")
            else:
                if save_patient_data(nom, email, reponses):
                    st.success("✅ Vos réponses ont été transmises."); st.balloons()

# ==============================================================================
# 2. ESPACE THÉRAPEUTE (ADMIN) - VERSION AVEC MÉMOIRE DE SESSION
# ==============================================================================
elif mode == "Espace Thérapeute":
    st.sidebar.divider()
    pwd_input = st.sidebar.text_input("Mot de passe Admin", type="password")
    
    if pwd_input == st.secrets["ADMIN_PASSWORD"]:
        st.header("🔒 Tableau de Bord Clinique Expert")
        df = load_all_patients()
        
        if df.empty:
            st.info("Aucun dossier pour le moment.")
        else:
            st.markdown("### Gestion des Dossiers")
            st.dataframe(df[["created_at", "nom", "email"]], use_container_width=True)
            st.divider()
            
            c_select, c_action = st.columns([3, 1])
            with c_select:
                patient_options = {f"{row['nom']} ({row['created_at'][:16]})": row['id'] for index, row in df.iterrows()}
                selected_label = st.selectbox("Sélectionner un dossier :", list(patient_options.keys()))
                selected_id = patient_options[selected_label]
            
            with c_action:
                st.write(""); st.write("")
                if st.button("🗑️ Supprimer", type="primary"):
                    if delete_patient(selected_id):
                        st.success("Dossier supprimé.")
                        st.rerun()
            
            st.markdown("---")
            
            # --- GESTION DE LA MÉMOIRE (SESSION STATE) ---
            # Si on change de patient, on réinitialise l'analyse
            if 'last_selected_id' not in st.session_state:
                st.session_state.last_selected_id = selected_id
            
            if st.session_state.last_selected_id != selected_id:
                st.session_state.analyse_active = False # On ferme l'analyse précédente
                st.session_state.last_selected_id = selected_id

            # Initialisation de la variable de mémoire
            if 'analyse_active' not in st.session_state:
                st.session_state.analyse_active = False

            # Récupération des données
            pat_data = df[df["id"] == selected_id].iloc[0]
            reponses_dict = decrypt_reponses(pat_data["reponses_json"])

            # 1. FICHIER BRUT
            col_raw, col_btn = st.columns(2)
            
            def gen_raw():
                doc = Document()
                doc.add_heading(f"Réponses Brutes : {pat_data['nom']}", 0)
                for i, (dom, q_d) in enumerate(YSQ_QUESTIONS.items()):
                    doc.add_heading(dom, 2)
                    for q, t in q_d.items():
                        s = reponses_dict.get(f"Q{q}", "-")
                        p = doc.add_paragraph(); p.add_run(f"Q{q}. ").bold=True; p.add_run(f"{t} : [{s}/6]")
                        if s != "-" and str(s).isdigit() and int(s) >= 5: p.runs[-1].font.color.rgb = RGBColor(255, 0, 0)
                out = BytesIO(); doc.save(out); return out.getvalue()
            
            with col_raw:
                st.download_button("📄 Télécharger Réponses Brutes", gen_raw(), f"Reponses_{pat_data['nom']}.docx")

            # 2. BOUTON D'ANALYSE (Active la mémoire)
            with col_btn:
                if st.button("📊 Lancer l'Analyse Clinique"):
                    st.session_state.analyse_active = True # On active la mémoire !

            # 3. AFFICHAGE PERSISTANT (Si la mémoire est active)
            if st.session_state.analyse_active:
                st.divider()
                
                # --- CALCULS ---
                resultats = []
                pre_selection = []
                
                for domaine, q_dict in YSQ_QUESTIONS.items():
                    code = domaine.split(" : ")[0]
                    nom_sch = domaine.split(" : ")[1]
                    scores = [int(reponses_dict.get(f"Q{k}", 1) or 1) for k in q_dict.keys()]
                    
                    moy = sum(scores) / len(scores)
                    sev = len([x for x in scores if x >= 5])
                    pct = (sev / len(scores)) * 100
                    etoile = " ⭐" if sev > 0 else ""
                    
                    # Logique de pré-sélection pour le rapport
                    if moy > 3.5:
                        niv = "🔴 IMPORTANT"
                        pre_selection.append(code)
                    elif moy >= 2.5:
                        niv = "🟡 Moyen"
                        pre_selection.append(code)
                    else:
                        niv = "🟢 Faible"
                        if sev > 0: pre_selection.append(code) # Ajoute si faible mais avec pics

                    resultats.append({
                        "Code": code,
                        "Schéma": f"{nom_sch}{etoile}",
                        "Moyenne": round(moy, 2),
                        "% Sévérité": f"{round(pct, 1)}%",
                        "Niveau": niv
                    })
                
                df_res = pd.DataFrame(resultats)
                
                # --- VISUALISATION ---
                c1, c2 = st.columns(2)
                with c1: st.table(df_res)
                with c2:
                    # Radar
                    fig_radar = px.line_polar(df_res, r='Moyenne', theta='Code', line_close=True, range_r=[0,6])
                    fig_radar.update_traces(fill='toself', line_color='blue')
                    st.plotly_chart(fig_radar)
                    
                    # Barres Dégradées
                    fig_bar = px.bar(df_res, x='Code', y='Moyenne', range_y=[0,6], 
                                     color="Moyenne", 
                                     color_continuous_scale="RdYlGn_r", 
                                     title="Intensité des Schémas")
                    st.plotly_chart(fig_bar)

                # --- SÉLECTION DES SCHÉMAS (PERSISTANTE MAINTENANT) ---
                st.markdown("---")
                st.subheader("📝 Personnalisation du Rapport")
                
                # La liste des codes disponibles
                codes_possibles = df_res["Code"].tolist()
                
                # Le Multiselect (ne disparaîtra plus !)
                selection_finale = st.multiselect(
                    "Ajoutez ou retirez des schémas pour le rapport Word :",
                    options=codes_possibles,
                    default=pre_selection
                )

                # --- GÉNÉRATION WORD ---
                def gen_expert():
                    doc = Document()
                    doc.add_heading(f"Bilan Psychométrique : {pat_data['nom']}", 0)
                    doc.add_paragraph(f"Date : {pat_data['created_at'][:10]}")
                    
                    doc.add_heading('1. Synthèse Visuelle', 1)
                    try: 
                        doc.add_picture(BytesIO(fig_radar.to_image(format="png", engine="kaleido")), width=Inches(4.5))
                        doc.add_picture(BytesIO(fig_bar.to_image(format="png", engine="kaleido")), width=Inches(4.5))
                    except: doc.add_paragraph("[Graphiques indisponibles]")

                    doc.add_heading('2. Tableau des Scores', 1)
                    tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'
                    h = tbl.rows[0].cells; h[0].text="Code"; h[1].text="Schéma"; h[2].text="Score"; h[3].text="Niveau"
                    for _, r in df_res.iterrows():
                        row = tbl.add_row().cells
                        row[0].text = str(r['Code']); row[1].text = str(r['Schéma']); row[2].text = str(r['Moyenne'])
                        run = row[3].paragraphs[0].add_run(r['Niveau']); run.bold = True
                        if "IMPORTANT" in r['Niveau']: run.font.color.rgb = RGBColor(255, 0, 0)
                        elif "Moyen" in r['Niveau']: run.font.color.rgb = RGBColor(255, 140, 0)
                        else: run.font.color.rgb = RGBColor(0, 128, 0)

                    doc.add_heading('3. Analyse Intégrale', 1)
                    
                    # Utilisation de la sélection manuelle
                    if selection_finale:
                        for d_name, d_info in YOUNG_DOMAINS_INFO.items():
                            # Filtre basé sur la sélection de l'utilisateur
                            match = [c for c in d_info["codes"] if c in selection_finale]
                            
                            if match:
                                doc.add_heading(d_name, 2)
                                doc.add_paragraph(d_info["besoin"]).italic = True
                                for c in match:
                                    inf = DATA_SCHEMAS.get(c, {})
                                    if not inf: continue
                                    
                                    p = doc.add_paragraph(); p.add_run(f"\n🔹 {inf['titre']}").bold = True
                                    p.add_run(f" - {inf['slogan']}").italic = True
                                    p.add_run(f" (Score: {df_res.loc[df_res['Code'] == c, 'Moyenne'].values[0]})")
                                    
                                    doc.add_paragraph("🧠 Analyse Clinique (Expert) :").bold = True
                                    doc.add_paragraph(inf['clinique_expert'])
                                    
                                    doc.add_paragraph("✝️ Perspective Théologique :").bold = True
                                    doc.add_paragraph(inf['theologie_expert'])
                                    
                                    doc.add_paragraph("🌱 Origines & Développement :").bold = True
                                    for o in inf['origines']: doc.add_paragraph(f"- {o}", style='List Bullet')
                                    
                                    doc.add_paragraph("⚠️ Signes au Quotidien :").bold = True
                                    for s in inf['symptomes']: doc.add_paragraph(f"- {s}", style='List Bullet')
                                    
                                    doc.add_paragraph(f"⚙️ Mécanisme Clé : {inf['mecanisme_titre']}").bold = True
                                    doc.add_paragraph(inf['mecanisme_texte'])
                                    
                                    doc.add_paragraph("👉 Plan d'Action Intégratif :").bold = True
                                    doc.add_paragraph("🛠️ Stratégies Thérapeutiques :").italic = True
                                    for act in inf['actions_therapeute']: doc.add_paragraph(f"• {act}")
                                    
                                    doc.add_paragraph("🙏 Conseil Pastoral :").italic = True
                                    doc.add_paragraph(inf['action_pastorale'])
                                    
                                    p_v = doc.add_paragraph(); p_v.add_run("📖 Verset d'ancrage : ").bold = True
                                    p_v.add_run(inf['verset']).italic = True
                                    doc.add_paragraph("-" * 30)
                    else: doc.add_paragraph("Aucun schéma sélectionné.")
                    
                    out = BytesIO(); doc.save(out); return out.getvalue()

                st.download_button("📥 Télécharger Rapport Expert", gen_expert(), f"Bilan_{pat_data['nom']}.docx")
    
    elif pwd_input: st.error("Mot de passe incorrect.")
# --- FOOTER ---
st.markdown("---")
with st.expander("🔒 Confidentialité et Protection des Données"):
    st.markdown("""
    **Engagement de confidentialité :**
    * Les données recueillies via ce formulaire sont strictement confidentielles.
    * Seul votre thérapeute a accès aux résultats détaillés.
    * Conformément à la loi, vous pouvez demander à tout moment la suppression intégrale de vos réponses.
    """)
    st.caption("Outil clinique YSQ-L3. Usage professionnel uniquement pour relation d'aide chrétienne- La Barque 2026.")
